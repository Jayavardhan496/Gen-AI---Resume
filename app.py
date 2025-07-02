import streamlit as st
import groq
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import io
import base64
from datetime import datetime

# Page configuration
st.set_page_config(
    page_title="AI Resume & Cover Letter Writer",
    page_icon="📄",
    layout="wide"
)

# Initialize Groq client
@st.cache_resource
def get_groq_client():
    try:
        client = groq.Groq(api_key=st.secrets["GROQ_API_KEY"])
        return client
    except Exception as e:
        st.error(f"Error initializing Groq client: {e}")
        return None

def generate_resume_content(client, user_data):
    """Generate resume content using Groq API"""
    prompt = f"""
    Create a professional resume for the following person. Structure it with clear sections and bullet points:

    Personal Information:
    - Name: {user_data['name']}
    - Email: {user_data['email']} 
    - Phone: {user_data['phone']}
    - Address: {user_data['address']}
    - LinkedIn: {user_data.get('linkedin', 'N/A')}

    Education: {user_data['education']}
    
    Work Experience: {user_data['work_experience']}
    
    Skills: {user_data['skills']}
    
    Target Job Role: {user_data['job_role']}
    
    Additional Information: {user_data.get('additional_info', 'None')}

    Please create a professional, ATS-friendly resume. Return ONLY the structured content in this exact format:

    PROFESSIONAL SUMMARY
    [Write a compelling 2-3 line professional summary highlighting key qualifications for the target role]

    EDUCATION
    [Format education details clearly]

    WORK EXPERIENCE  
    [Format work experience with company names, positions, and bullet points of achievements]

    SKILLS
    [Organize skills by category if applicable]

    ADDITIONAL INFORMATION
    [Include certifications, projects, awards if provided]

    Use clear formatting with bullet points and ensure each section is well-organized.
    """
    
    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-8b-8192",
            temperature=0.7,
            max_tokens=1500
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error generating resume: {e}")
        return None

def generate_cover_letter(client, user_data, company_info):
    """Generate cover letter using Groq API"""
    prompt = f"""
    Write a professional cover letter for the following person applying to {company_info.get('company_name', 'the company')} for the position of {user_data['job_role']}.

    Applicant Information:
    - Name: {user_data['name']}
    - Target Role: {user_data['job_role']}
    - Work Experience: {user_data['work_experience']}
    - Skills: {user_data['skills']}
    - Education: {user_data['education']}

    Company Information:
    - Company Name: {company_info.get('company_name', 'Company')}
    - Job Description/Requirements: {company_info.get('job_description', 'Not provided')}
    - Hiring Manager: {company_info.get('hiring_manager', 'Hiring Manager')}

    Create a compelling, personalized cover letter that:
    1. Shows enthusiasm for the role and company
    2. Highlights relevant experience and skills
    3. Demonstrates knowledge of the company/role
    4. Has a professional tone
    5. Is concise but impactful
    
    Format it as a proper business letter.
    """
    
    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-8b-8192",
            temperature=0.7,
            max_tokens=1200
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error generating cover letter: {e}")
        return None

def create_resume_html(resume_content, user_data):
    """Convert resume content to HTML with styling"""
    
    # Parse the resume content into sections
    sections = parse_resume_sections(resume_content)
    
    html_sections = ""
    for section_title, section_content in sections.items():
        formatted_content = format_section_content(section_content)
        html_sections += f"""
            <div class="section">
                <div class="section-title">{section_title}</div>
                <div class="section-content">{formatted_content}</div>
            </div>
        """
    
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: 'Arial', sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                background-color: #f9f9f9;
            }}
            .resume-container {{
                background: white;
                padding: 40px;
                border-radius: 8px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }}
            .header {{
                text-align: center;
                border-bottom: 3px solid #2c3e50;
                padding-bottom: 20px;
                margin-bottom: 30px;
            }}
            .name {{
                font-size: 32px;
                font-weight: bold;
                color: #2c3e50;
                margin-bottom: 10px;
                letter-spacing: 1px;
            }}
            .contact-info {{
                font-size: 14px;
                color: #666;
                line-height: 1.4;
            }}
            .section {{
                margin-bottom: 25px;
            }}
            .section-title {{
                font-size: 16px;
                font-weight: bold;
                color: #2c3e50;
                border-bottom: 2px solid #3498db;
                padding-bottom: 5px;
                margin-bottom: 15px;
                text-transform: uppercase;
                letter-spacing: 0.5px;
            }}
            .section-content {{
                font-size: 14px;
                line-height: 1.6;
            }}
            .section-content ul {{
                margin: 10px 0;
                padding-left: 20px;
            }}
            .section-content li {{
                margin-bottom: 8px;
                text-align: justify;
            }}
            .section-content p {{
                margin-bottom: 12px;
                text-align: justify;
            }}
            .work-item {{
                margin-bottom: 20px;
            }}
            .job-title {{
                font-weight: bold;
                color: #2c3e50;
                font-size: 15px;
            }}
            .company {{
                color: #3498db;
                font-weight: 600;
            }}
            .date {{
                color: #666;
                font-style: italic;
                float: right;
            }}
            @media print {{
                body {{ background-color: white; }}
                .resume-container {{ box-shadow: none; }}
            }}
        </style>
    </head>
    <body>
        <div class="resume-container">
            <div class="header">
                <div class="name">{user_data['name']}</div>
                <div class="contact-info">
                    {user_data['email']} | {user_data['phone']}<br>
                    {user_data['address']}
                    {f"<br>LinkedIn: {user_data['linkedin']}" if user_data.get('linkedin') and user_data['linkedin'].strip() else ""}
                </div>
            </div>
            {html_sections}
        </div>
    </body>
    </html>
    """
    return html_content

def parse_resume_sections(resume_content):
    """Parse resume content into sections"""
    sections = {}
    current_section = None
    current_content = []
    
    lines = resume_content.split('\n')
    
    section_keywords = [
        'PROFESSIONAL SUMMARY', 'SUMMARY', 'PROFILE',
        'EDUCATION', 'ACADEMIC BACKGROUND',
        'WORK EXPERIENCE', 'EXPERIENCE', 'EMPLOYMENT', 'PROFESSIONAL EXPERIENCE',
        'SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES',
        'ADDITIONAL INFORMATION', 'CERTIFICATIONS', 'PROJECTS', 'ACHIEVEMENTS'
    ]
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check if this line is a section header
        is_section_header = False
        for keyword in section_keywords:
            if keyword in line.upper() and len(line) <= 50:  # Likely a header
                is_section_header = True
                if current_section:
                    sections[current_section] = '\n'.join(current_content)
                current_section = keyword
                current_content = []
                break
        
        if not is_section_header and current_section:
            current_content.append(line)
    
    # Add the last section
    if current_section and current_content:
        sections[current_section] = '\n'.join(current_content)
    
    return sections

def format_section_content(content):
    """Format section content with proper HTML"""
    if not content.strip():
        return "<p>Information not provided</p>"
    
    lines = content.split('\n')
    formatted_lines = []
    current_paragraph = []
    
    for line in lines:
        line = line.strip()
        if not line:
            if current_paragraph:
                formatted_lines.append('<p>' + ' '.join(current_paragraph) + '</p>')
                current_paragraph = []
            continue
            
        # Check if line starts with bullet point indicators
        if line.startswith(('•', '-', '*', '▪', '◦')) or line.startswith(tuple(f'{i}.' for i in range(1, 10))):
            if current_paragraph:
                formatted_lines.append('<p>' + ' '.join(current_paragraph) + '</p>')
                current_paragraph = []
            # Start or continue a list
            if not formatted_lines or not formatted_lines[-1].endswith('</ul>'):
                formatted_lines.append('<ul>')
            # Clean the bullet point
            clean_line = line[1:].strip() if line[0] in '•-*▪◦' else line[2:].strip()
            formatted_lines.append(f'<li>{clean_line}</li>')
        else:
            # Close any open list
            if formatted_lines and formatted_lines[-1].startswith('<li>'):
                formatted_lines.append('</ul>')
            current_paragraph.append(line)
    
    # Handle remaining content
    if current_paragraph:
        formatted_lines.append('<p>' + ' '.join(current_paragraph) + '</p>')
    
    # Close any unclosed list
    if formatted_lines and formatted_lines[-1].startswith('<li>'):
        formatted_lines.append('</ul>')
    
    return ''.join(formatted_lines)

def create_cover_letter_html(cover_letter_content, user_data):
    """Convert cover letter to HTML with styling"""
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{
                font-family: 'Arial', sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                background-color: #f9f9f9;
            }}
            .letter-container {{
                background: white;
                padding: 40px;
                border-radius: 8px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            }}
            .header {{
                text-align: right;
                margin-bottom: 30px;
            }}
            .sender-info {{
                font-size: 14px;
                color: #666;
                line-height: 1.4;
            }}
            .date {{
                margin: 20px 0;
                font-size: 14px;
            }}
            .content {{
                white-space: pre-line;
                text-align: justify;
                font-size: 14px;
                line-height: 1.6;
            }}
            @media print {{
                body {{ background-color: white; }}
                .letter-container {{ box-shadow: none; }}
            }}
        </style>
    </head>
    <body>
        <div class="letter-container">
            <div class="header">
                <div class="sender-info">
                    {user_data['name']}<br>
                    {user_data['email']}<br>
                    {user_data['phone']}<br>
                    {user_data['address']}
                </div>
            </div>
            <div class="date">{datetime.now().strftime('%B %d, %Y')}</div>
            <div class="content">{cover_letter_content}</div>
        </div>
    </body>
    </html>
    """
    return html_content

def create_docx_resume(resume_content, user_data):
    """Create a Word document for the resume"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title - Name
    title = doc.add_heading(user_data['name'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact_para = doc.add_paragraph()
    contact_run = contact_para.add_run(f"{user_data['email']} | {user_data['phone']}")
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    address_para = doc.add_paragraph()
    address_run = address_para.add_run(user_data['address'])
    address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if user_data.get('linkedin') and user_data['linkedin'].strip():
        linkedin_para = doc.add_paragraph()
        linkedin_run = linkedin_para.add_run(f"LinkedIn: {user_data['linkedin']}")
        linkedin_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add line break
    doc.add_paragraph()
    
    # Parse and add resume sections
    sections_dict = parse_resume_sections(resume_content)
    
    for section_title, section_content in sections_dict.items():
        # Add section heading
        heading = doc.add_heading(section_title, level=1)
        
        # Add section content
        lines = section_content.split('\n')
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    para = doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                continue
                
            # Check if line is a bullet point
            if line.startswith(('•', '-', '*', '▪', '◦')) or line.startswith(tuple(f'{i}.' for i in range(1, 10))):
                if current_paragraph:
                    para = doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                # Add bullet point
                clean_line = line[1:].strip() if line[0] in '•-*▪◦' else line[2:].strip()
                bullet_para = doc.add_paragraph(clean_line, style='List Bullet')
            else:
                current_paragraph.append(line)
        
        # Add any remaining paragraph content
        if current_paragraph:
            para = doc.add_paragraph(' '.join(current_paragraph))
        
        # Add space after each section
        doc.add_paragraph()
    
    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def create_docx_cover_letter(cover_letter_content, user_data):
    """Create a Word document for the cover letter"""
    doc = Document()
    
    # Header with contact info
    header_para = doc.add_paragraph()
    header_para.add_run(f"{user_data['name']}\n{user_data['email']}\n{user_data['phone']}\n{user_data['address']}")
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add date
    date_para = doc.add_paragraph()
    date_para.add_run(datetime.now().strftime('%B %d, %Y'))
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add a line break
    doc.add_paragraph()
    
    # Cover letter content
    doc.add_paragraph(cover_letter_content)
    
    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def main():
    st.title("🎯 AI Resume & Cover Letter Writer")
    st.markdown("Create professional resumes and cover letters tailored to your target job role.")
    
    # Initialize Groq client
    client = get_groq_client()
    if not client:
        st.stop()
    
    # Sidebar for user inputs
    st.sidebar.header("📋 Your Information")
    
    # Personal Information
    st.sidebar.subheader("Personal Details")
    name = st.sidebar.text_input("Full Name*", placeholder="John Doe")
    email = st.sidebar.text_input("Email*", placeholder="john.doe@email.com")
    phone = st.sidebar.text_input("Phone*", placeholder="+1 (555) 123-4567")
    address = st.sidebar.text_area("Address*", placeholder="123 Main St, City, State 12345")
    linkedin = st.sidebar.text_input("LinkedIn Profile (optional)", placeholder="https://linkedin.com/in/johndoe")
    
    # Professional Information
    st.sidebar.subheader("Professional Information")
    job_role = st.sidebar.text_input("Target Job Role*", placeholder="Software Developer")
    education = st.sidebar.text_area("Education*", placeholder="Bachelor's in Computer Science, XYZ University (2020)")
    work_experience = st.sidebar.text_area("Work Experience*", placeholder="Software Engineer at ABC Corp (2020-2023)\n- Developed web applications\n- Led team of 3 developers")
    skills = st.sidebar.text_area("Skills*", placeholder="Python, JavaScript, React, SQL, Git")
    additional_info = st.sidebar.text_area("Additional Information (optional)", placeholder="Certifications, Projects, Awards, etc.")
    
    # Company-specific information for cover letter
    st.sidebar.subheader("Cover Letter Details")
    company_name = st.sidebar.text_input("Company Name", placeholder="Tech Corp Inc.")
    hiring_manager = st.sidebar.text_input("Hiring Manager Name (optional)", placeholder="Ms. Jane Smith")
    job_description = st.sidebar.text_area("Job Description/Requirements (optional)", placeholder="Paste job posting details here...")
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📄 Resume")
        if st.button("Generate Resume", type="primary", use_container_width=True):
            if not all([name, email, phone, address, job_role, education, work_experience, skills]):
                st.error("Please fill in all required fields marked with *")
            else:
                with st.spinner("Generating your professional resume..."):
                    user_data = {
                        'name': name,
                        'email': email,
                        'phone': phone,
                        'address': address,
                        'linkedin': linkedin,
                        'job_role': job_role,
                        'education': education,
                        'work_experience': work_experience,
                        'skills': skills,
                        'additional_info': additional_info
                    }
                    
                    resume_content = generate_resume_content(client, user_data)
                    if resume_content:
                        st.session_state['resume_content'] = resume_content
                        st.session_state['user_data'] = user_data
                        st.success("Resume generated successfully!")
        
        # Display resume if generated
        if 'resume_content' in st.session_state:
            st.subheader("📋 Resume Preview")
            resume_html = create_resume_html(st.session_state['resume_content'], st.session_state['user_data'])
            st.components.v1.html(resume_html, height=800, scrolling=True)
            
            # Download button for resume
            resume_docx = create_docx_resume(st.session_state['resume_content'], st.session_state['user_data'])
            st.download_button(
                label="📥 Download Resume (.docx)",
                data=resume_docx.getvalue(),
                file_name=f"{st.session_state['user_data']['name']}_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    with col2:
        st.header("✉️ Cover Letter")
        if st.button("Generate Cover Letter", type="primary", use_container_width=True):
            if not all([name, email, phone, address, job_role, education, work_experience, skills]):
                st.error("Please fill in all required fields marked with *")
            elif not company_name:
                st.error("Please enter the company name for the cover letter")
            else:
                with st.spinner("Crafting your personalized cover letter..."):
                    user_data = {
                        'name': name,
                        'email': email,
                        'phone': phone,
                        'address': address,
                        'linkedin': linkedin,
                        'job_role': job_role,
                        'education': education,
                        'work_experience': work_experience,
                        'skills': skills,
                        'additional_info': additional_info
                    }
                    
                    company_info = {
                        'company_name': company_name,
                        'hiring_manager': hiring_manager,
                        'job_description': job_description
                    }
                    
                    cover_letter_content = generate_cover_letter(client, user_data, company_info)
                    if cover_letter_content:
                        st.session_state['cover_letter_content'] = cover_letter_content
                        st.session_state['user_data'] = user_data
                        st.success("Cover letter generated successfully!")
        
        # Display cover letter if generated
        if 'cover_letter_content' in st.session_state:
            st.subheader("📝 Cover Letter Preview")
            cover_letter_html = create_cover_letter_html(st.session_state['cover_letter_content'], st.session_state['user_data'])
            st.components.v1.html(cover_letter_html, height=800, scrolling=True)
            
            # Download button for cover letter
            cover_letter_docx = create_docx_cover_letter(st.session_state['cover_letter_content'], st.session_state['user_data'])
            st.download_button(
                label="📥 Download Cover Letter (.docx)",
                data=cover_letter_docx.getvalue(),
                file_name=f"{st.session_state['user_data']['name']}_Cover_Letter.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    
    # Footer
    st.markdown("---")
    st.markdown("💡 **Tips for better results:**")
    st.markdown("- Provide detailed work experience with specific achievements")
    st.markdown("- Include relevant skills for your target job role")
    st.markdown("- Add company-specific information for personalized cover letters")
    st.markdown("- Review and customize the generated content before using")

if __name__ == "__main__":
    main()
